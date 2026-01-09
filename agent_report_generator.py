"""
Agent 6: Report Generator
Responsible for: PRISMA flow diagram generation, report writing,
table/figure creation, and multi-format export.
"""
from typing import Dict, Any, List, Optional
import logging
from datetime import datetime
from pathlib import Path
from collections import Counter
from config import settings
from state import (
    Document, PRISMAFlowDiagram, ReviewReport,
    SynthesisResult, PRISMAState, add_audit_entry
)

logger = logging.getLogger(__name__)


# ============================================================================
# REPORT GENERATOR AGENT
# ============================================================================

class ReportGeneratorAgent:
    """
    Agent 6: Report Generator
    
    Responsibilities:
    1. Generate PRISMA 2020 flow diagram
    2. Create tables and figures
    3. Write structured report sections
    4. Format citations
    5. Export to multiple formats (PDF, DOCX, HTML)
    """
    
    def __init__(self):
        self.output_dir = Path(settings.document_store_path) / "reports"
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_prisma_flow_diagram(
        self,
        retrieval_stats: Dict[str, Any],
        screening_decisions: List[Any],
        included_docs: List[Document],
        excluded_docs: List[Document],
        borderline_docs: List[Document]
    ) -> PRISMAFlowDiagram:
        """
        Generate PRISMA 2020 flow diagram data.
        
        Returns:
            PRISMAFlowDiagram with all counts
        """
        logger.info("Generating PRISMA flow diagram...")
        
        # Identification phase
        total_retrieved = retrieval_stats.get("total_retrieved", 0)
        duplicates_removed = retrieval_stats.get("duplicates_removed", 0)
        records_screened = total_retrieved - duplicates_removed
        
        # Screening phase
        excluded_count = len(excluded_docs)
        
        # Eligibility phase (in a full implementation, would track full-text screening)
        # For now, borderline papers represent those needing full-text review
        full_text_assessed = len(included_docs) + len(borderline_docs)
        full_text_excluded = len(borderline_docs)  # Simplified
        
        # Included phase
        included_in_synthesis = len(included_docs)
        
        flow = PRISMAFlowDiagram(
            identification={
                "records_identified_via_databases": total_retrieved,
                "records_from_other_sources": 0,
                "total_records": total_retrieved,
                "duplicates_removed": duplicates_removed,
                "records_screened": records_screened
            },
            screening={
                "records_screened": records_screened,
                "records_excluded": excluded_count,
                "records_assessed_for_eligibility": full_text_assessed
            },
            eligibility={
                "full_text_articles_assessed": full_text_assessed,
                "full_text_articles_excluded": full_text_excluded,
                "excluded_reasons": self._get_exclusion_reasons(excluded_docs, screening_decisions)
            },
            included={
                "studies_included_in_synthesis": included_in_synthesis,
                "studies_included_in_meta_analysis": 0  # N/A for this implementation
            }
        )
        
        return flow
    
    def _get_exclusion_reasons(
        self,
        excluded_docs: List[Document],
        screening_decisions: List[Any]
    ) -> Dict[str, int]:
        """Extract and count exclusion reasons."""
        reasons = []
        
        excluded_ids = {doc.id for doc in excluded_docs}
        for decision in screening_decisions:
            if decision.document_id in excluded_ids:
                # Use first violated criterion as primary reason
                if decision.criteria_violated:
                    reasons.append(decision.criteria_violated[0])
        
        return dict(Counter(reasons))
    
    def create_prisma_flow_diagram_image(
        self,
        flow: PRISMAFlowDiagram,
        output_path: Optional[str] = None
    ) -> str:
        """
        Create visual PRISMA flow diagram.
        
        Returns:
            Path to generated diagram image
        """
        try:
            import matplotlib.pyplot as plt
            import matplotlib.patches as mpatches
            from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
            
            if output_path is None:
                output_path = self.output_dir / "prisma_flow_diagram.png"
            
            fig, ax = plt.subplots(figsize=(12, 14))
            ax.set_xlim(0, 10)
            ax.set_ylim(0, 20)
            ax.axis('off')
            
            # Box styling
            box_props = dict(boxstyle="round,pad=0.3", facecolor='lightblue', edgecolor='black', linewidth=2)
            
            # Title
            ax.text(5, 19, 'PRISMA 2020 Flow Diagram', ha='center', va='top',
                   fontsize=16, fontweight='bold')
            
            # Identification
            y_pos = 17
            ax.text(5, y_pos, f"Records identified via databases\n(n = {flow.identification['records_identified_via_databases']})",
                   ha='center', va='center', bbox=box_props, fontsize=10)
            
            # Arrow down
            ax.annotate('', xy=(5, y_pos - 1), xytext=(5, y_pos - 0.5),
                       arrowprops=dict(arrowstyle='->', lw=2))
            
            # Duplicates removed
            y_pos = 15
            ax.text(5, y_pos, f"Records after duplicates removed\n(n = {flow.identification['records_screened']})",
                   ha='center', va='center', bbox=box_props, fontsize=10)
            
            # Duplicates removed box (side)
            ax.text(8.5, y_pos, f"Duplicates removed\n(n = {flow.identification['duplicates_removed']})",
                   ha='center', va='center', bbox=dict(boxstyle="round,pad=0.3", facecolor='lightcoral'),
                   fontsize=9)
            
            # Arrow down
            ax.annotate('', xy=(5, y_pos - 1.5), xytext=(5, y_pos - 0.5),
                       arrowprops=dict(arrowstyle='->', lw=2))
            
            # Screening
            y_pos = 12.5
            ax.text(5, y_pos, f"Records screened\n(n = {flow.screening['records_screened']})",
                   ha='center', va='center', bbox=box_props, fontsize=10)
            
            # Records excluded (side)
            ax.text(8.5, y_pos, f"Records excluded\n(n = {flow.screening['records_excluded']})",
                   ha='center', va='center', bbox=dict(boxstyle="round,pad=0.3", facecolor='lightcoral'),
                   fontsize=9)
            
            # Arrow down
            ax.annotate('', xy=(5, y_pos - 1.5), xytext=(5, y_pos - 0.5),
                       arrowprops=dict(arrowstyle='->', lw=2))
            
            # Eligibility
            y_pos = 10
            ax.text(5, y_pos, f"Full-text articles assessed\n(n = {flow.eligibility['full_text_articles_assessed']})",
                   ha='center', va='center', bbox=box_props, fontsize=10)
            
            # Full-text excluded (side) with reasons
            exclusion_text = f"Full-text excluded\n(n = {flow.eligibility['full_text_articles_excluded']})"
            ax.text(8.5, y_pos, exclusion_text,
                   ha='center', va='center', bbox=dict(boxstyle="round,pad=0.3", facecolor='lightcoral'),
                   fontsize=9)
            
            # Arrow down
            ax.annotate('', xy=(5, y_pos - 1.5), xytext=(5, y_pos - 0.5),
                       arrowprops=dict(arrowstyle='->', lw=2))
            
            # Included
            y_pos = 7.5
            ax.text(5, y_pos, f"Studies included in synthesis\n(n = {flow.included['studies_included_in_synthesis']})",
                   ha='center', va='center', bbox=dict(boxstyle="round,pad=0.3", facecolor='lightgreen', linewidth=2),
                   fontsize=11, fontweight='bold')
            
            plt.tight_layout()
            plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()
            
            logger.info(f"PRISMA flow diagram saved to {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to create flow diagram image: {e}")
            return ""
    
    def generate_report_text(
        self,
        research_question: str,
        pico: Any,
        search_queries: List[Any],
        screening_criteria: Any,
        included_docs: List[Document],
        synthesis: SynthesisResult,
        prisma_flow: PRISMAFlowDiagram
    ) -> str:
        """Generate structured report text."""
        report = []
        
        # Title
        report.append("# Systematic Review Report\n")
        report.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        report.append(f"**Research Question:** {research_question}\n\n")
        
        # DISCLAIMER
        report.append("---\n")
        report.append("**⚠️ IMPORTANT DISCLAIMER**\n\n")
        report.append("This report was generated by an automated AI-powered system (AutoPRISMA) for demonstration purposes only. ")
        report.append("All findings require expert validation and should NOT be used for publication, clinical decisions, or policy-making ")
        report.append("without thorough human review by qualified domain experts.\n")
        report.append("---\n\n")
        
        # 1. Introduction
        report.append("## 1. Introduction\n\n")
        report.append(f"This systematic review examined: {research_question}\n\n")
        
        # 2. Methods
        report.append("## 2. Methods\n\n")
        
        # 2.1 PICO
        report.append("### 2.1 PICO Framework\n\n")
        report.append(f"- **Population:** {pico.population}\n")
        report.append(f"- **Intervention:** {pico.intervention}\n")
        report.append(f"- **Comparator:** {pico.comparator}\n")
        report.append(f"- **Outcome:** {pico.outcome}\n")
        report.append(f"- **Study Types:** {', '.join(pico.study_types)}\n\n")
        
        # 2.2 Search Strategy
        report.append("### 2.2 Search Strategy\n\n")
        report.append(f"Searches were conducted across {len(set(db for q in search_queries for db in q.databases))} databases.\n\n")
        for i, query in enumerate(search_queries, 1):
            report.append(f"**Query {i}:** {query.boolean_query}\n\n")
        
        # 2.3 Screening Criteria
        report.append("### 2.3 Screening Criteria\n\n")
        report.append("**Inclusion Criteria:**\n")
        for criterion in screening_criteria.inclusion_criteria:
            report.append(f"- {criterion}\n")
        report.append("\n**Exclusion Criteria:**\n")
        for criterion in screening_criteria.exclusion_criteria:
            report.append(f"- {criterion}\n")
        report.append("\n")
        
        # 3. Results
        report.append("## 3. Results\n\n")
        
        # 3.1 Study Selection
        report.append("### 3.1 Study Selection\n\n")
        report.append(f"The search identified {prisma_flow.identification['total_records']} records. ")
        report.append(f"After removing {prisma_flow.identification['duplicates_removed']} duplicates, ")
        report.append(f"{prisma_flow.identification['records_screened']} records were screened. ")
        report.append(f"Finally, {prisma_flow.included['studies_included_in_synthesis']} studies were included in the synthesis.\n\n")
        report.append("See PRISMA flow diagram for details.\n\n")
        
        # 3.2 Study Characteristics
        report.append("### 3.2 Study Characteristics\n\n")
        stats = synthesis.summary_statistics
        report.append(f"- **Total studies:** {stats['total_papers']}\n")
        if stats['year_range']:
            report.append(f"- **Year range:** {stats['year_range'][0]}-{stats['year_range'][1]}\n")
        report.append(f"- **Average authors per paper:** {stats['average_authors_per_paper']:.1f}\n\n")
        
        # 3.3 Synthesis of Findings
        report.append("### 3.3 Synthesis of Findings\n\n")
        
        # Themes
        report.append("**Major Themes:**\n\n")
        for theme in synthesis.themes:
            report.append(f"**{theme.name}** (*{theme.strength} evidence*)\n")
            report.append(f"{theme.description}\n\n")
        
        # Consensus
        if synthesis.consensus_findings:
            report.append("**Areas of Consensus:**\n\n")
            for finding in synthesis.consensus_findings:
                report.append(f"- {finding}\n")
            report.append("\n")
        
        # Contradictions
        if synthesis.contradictions:
            report.append("**Areas of Contradiction:**\n\n")
            for contradiction in synthesis.contradictions:
                report.append(f"- {contradiction}\n")
            report.append("\n")
        
        # 4. Discussion
        report.append("## 4. Discussion\n\n")
        
        # Research gaps
        report.append("### 4.1 Research Gaps\n\n")
        for gap in synthesis.research_gaps:
            report.append(f"- {gap}\n")
        report.append("\n")
        
        # Limitations
        report.append("### 4.2 Limitations\n\n")
        report.append("- This review was conducted using automated methods and requires expert validation\n")
        report.append("- Full-text review was not conducted for all papers\n")
        report.append("- Quality assessment of individual studies was not performed\n")
        report.append("- Meta-analysis was not conducted\n\n")
        
        # 5. Conclusion
        report.append("## 5. Conclusion\n\n")
        report.append(f"This systematic review synthesized {len(included_docs)} studies addressing {research_question.lower()}. ")
        report.append(f"The findings revealed {len(synthesis.themes)} major themes and identified {len(synthesis.research_gaps)} research gaps. ")
        report.append("Further research is needed in the areas identified.\n\n")
        
        # 6. References
        report.append("## 6. References\n\n")
        for i, doc in enumerate(included_docs, 1):
            authors_str = ", ".join([a.name for a in doc.authors[:3]])
            if len(doc.authors) > 3:
                authors_str += " et al."
            report.append(f"{i}. {authors_str} ({doc.year}). {doc.title}. *{doc.journal or 'Unknown Journal'}*.\n")
        
        return "".join(report)
    
    def save_report(
        self,
        report_text: str,
        filename: str = "systematic_review_report"
    ) -> Dict[str, str]:
        """
        Save report in multiple formats.
        
        Returns:
            Dictionary with format -> path mappings
        """
        paths = {}
        
        # Save as Markdown
        md_path = self.output_dir / f"{filename}.md"
        md_path.write_text(report_text, encoding="utf-8")
        paths["markdown"] = str(md_path)
        logger.info(f"Saved Markdown report to {md_path}")
        
        # Save as HTML
        try:
            import markdown
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>Systematic Review Report</title>
                <style>
                    body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
                    h1, h2, h3 {{ color: #333; }}
                    h1 {{ border-bottom: 2px solid #333; padding-bottom: 10px; }}
                    h2 {{ border-bottom: 1px solid #666; padding-bottom: 5px; margin-top: 30px; }}
                    hr {{ margin: 20px 0; }}
                </style>
            </head>
            <body>
            {markdown.markdown(report_text)}
            </body>
            </html>
            """
            html_path = self.output_dir / f"{filename}.html"
            html_path.write_text(html_content, encoding="utf-8")
            paths["html"] = str(html_path)
            logger.info(f"Saved HTML report to {html_path}")
        except ImportError:
            logger.warning("markdown package not available, skipping HTML export")
        
        # Save as DOCX
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = DocxDocument()
            
            # Parse and add content (simplified)
            lines = report_text.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)
            
            docx_path = self.output_dir / f"{filename}.docx"
            doc.save(str(docx_path))
            paths["docx"] = str(docx_path)
            logger.info(f"Saved DOCX report to {docx_path}")
        except ImportError:
            logger.warning("python-docx package not available, skipping DOCX export")
        
        return paths
    
    def run(self, state: PRISMAState) -> Dict[str, Any]:
        """
        Execute the Report Generator agent.
        
        Args:
            state: Current PRISMA workflow state
        
        Returns:
            State updates with final report
        """
        logger.info("=== Report Generator Agent Started ===")
        
        # Gather all necessary data
        research_question = state["research_question"]
        pico = state.get("pico_criteria")
        search_queries = state.get("search_queries", [])
        screening_criteria = state.get("screening_criteria")
        retrieval_stats = state.get("retrieval_stats", {})
        screening_decisions = state.get("screening_decisions", [])
        included_docs = state.get("included_documents", [])
        excluded_docs = state.get("excluded_documents", [])
        borderline_docs = state.get("borderline_documents", [])
        synthesis = state.get("synthesis_result")
        
        # Step 1: Generate PRISMA flow diagram
        logger.info("Step 1: Generating PRISMA flow diagram")
        prisma_flow = self.generate_prisma_flow_diagram(
            retrieval_stats,
            screening_decisions,
            included_docs,
            excluded_docs,
            borderline_docs
        )
        
        # Step 2: Create flow diagram image
        logger.info("Step 2: Creating flow diagram visualization")
        diagram_path = self.create_prisma_flow_diagram_image(prisma_flow)
        prisma_flow.diagram_path = diagram_path
        
        # Step 3: Generate report text
        logger.info("Step 3: Generating report text")
        report_text = self.generate_report_text(
            research_question,
            pico,
            search_queries,
            screening_criteria,
            included_docs,
            synthesis,
            prisma_flow
        )
        
        # Step 4: Save report in multiple formats
        logger.info("Step 4: Saving report in multiple formats")
        report_paths = self.save_report(report_text)
        
        # Create ReviewReport object
        exclusion_summary = {}
        for decision in screening_decisions:
            if decision.decision == "exclude" and decision.criteria_violated:
                reason = decision.criteria_violated[0]
                exclusion_summary[reason] = exclusion_summary.get(reason, 0) + 1
        
        final_report = ReviewReport(
            research_question=research_question,
            pico=pico,
            search_strategy=search_queries[0] if search_queries else None,
            screening_criteria=screening_criteria,
            prisma_flow=prisma_flow,
            synthesis=synthesis,
            included_papers=included_docs,
            excluded_papers_summary=exclusion_summary,
            report_text=report_text,
            report_path=report_paths.get("markdown"),
            metadata={
                "report_paths": report_paths,
                "diagram_path": diagram_path
            }
        )
        
        # Prepare state updates
        updates = {
            "prisma_flow": prisma_flow,
            "final_report": final_report,
            "current_stage": "report_generation_complete",
            "workflow_status": "completed",
            "completed_at": datetime.now()
        }
        
        # Add audit entry
        audit_entry = add_audit_entry(
            state,
            agent="ReportGenerator",
            action="generate_report",
            details={
                "report_paths": report_paths,
                "diagram_path": diagram_path,
                "sections_generated": ["introduction", "methods", "results", "discussion", "conclusion", "references"]
            }
        )
        updates.update(audit_entry)
        
        logger.info("=== Report Generator Agent Completed ===")
        return updates


# ============================================================================
# STANDALONE TESTING
# ============================================================================

if __name__ == "__main__":
    from state import create_initial_state, PICOCriteria, ScreeningCriteria, Author, Theme, SynthesisResult
    
    # Create minimal test state
    pico = PICOCriteria(
        population="Adults with anxiety",
        intervention="CBT",
        comparator="Standard care",
        outcome="Anxiety reduction",
        study_types=["RCT"]
    )
    
    test_docs = [
        Document(
            id=f"test{i}",
            title=f"Study {i}",
            authors=[Author(name=f"Author {i}")],
            abstract="Test abstract",
            year=2020,
            journal="Test Journal",
            source="test"
        )
        for i in range(5)
    ]
    
    synthesis = SynthesisResult(
        themes=[Theme(
            theme_id="t1",
            name="CBT Effectiveness",
            description="Studies show CBT is effective",
            supporting_papers=["test1", "test2"]
        )],
        research_gaps=["Need more long-term studies"],
        contradictions=[],
        consensus_findings=["CBT reduces anxiety"],
        summary_statistics={"total_papers": 5}
    )
    
    state = create_initial_state(research_question="CBT for anxiety")
    state["pico_criteria"] = pico
    state["included_documents"] = test_docs
    state["synthesis_result"] = synthesis
    state["retrieval_stats"] = {"total_retrieved": 100, "duplicates_removed": 20}
    state["screening_decisions"] = []
    state["excluded_documents"] = []
    state["borderline_documents"] = []
    state["screening_criteria"] = ScreeningCriteria(
        inclusion_criteria=["Adults", "CBT", "Anxiety"],
        exclusion_criteria=["Children", "No CBT"]
    )
    
    agent = ReportGeneratorAgent()
    result = agent.run(state)
    
    print("\n=== Report Generated ===")
    print(f"PRISMA Flow: {result['prisma_flow']}")
    print(f"Report paths: {result['final_report'].metadata['report_paths']}")
